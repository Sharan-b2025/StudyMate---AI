"""
File Processing Service
------------------------
Extracts raw text from uploaded study materials.
Supports: PDF, DOCX, TXT, PNG/JPG/JPEG (via OCR).

All extraction failures are caught and returned as (None, error_message)
so upload flows never crash on a corrupted/unsupported file.
"""
import os
import re
from flask import current_app

PAGE_MARKER_RE = re.compile(r"\[\[PAGE:(\d+)\]\]")
CHARS_PER_PSEUDO_PAGE = 3000  # used as a "page" size for DOCX/TXT/OCR text (no real pages)


class ExtractionError(Exception):
    pass


def allowed_file(filename):
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    return ext in current_app.config["ALLOWED_EXTENSIONS"], ext


def extract_text(filepath, file_type):
    try:
        if file_type == "pdf":
            return _extract_pdf(filepath), None
        if file_type == "docx":
            return _extract_docx(filepath), None
        if file_type == "txt":
            return _extract_txt(filepath), None
        if file_type in ("png", "jpg", "jpeg"):
            return _extract_image(filepath), None
        return None, f"Unsupported file type: {file_type}"
    except Exception as exc:  # noqa: BLE001 - we want to catch everything here
        return None, f"Could not process file: {exc}"


def _extract_pdf(filepath):
    from pypdf import PdfReader

    reader = PdfReader(filepath)
    pages = []
    for page in reader.pages:
        content = page.extract_text() or ""
        pages.append(content)
    text = "\n\n".join(f"[[PAGE:{i}]]\n{content}" for i, content in enumerate(pages, start=1)).strip()
    if not text.replace("[[PAGE:", "").strip():
        raise ExtractionError("No extractable text found (the PDF may be scanned/image-only).")
    return text


def _extract_docx(filepath):
    import docx

    document = docx.Document(filepath)
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    text = "\n".join(parts).strip()
    if not text:
        raise ExtractionError("The document appears to be empty.")
    return text


def _extract_txt(filepath):
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read().strip()
    if not text:
        raise ExtractionError("The text file is empty.")
    return text


def _extract_image(filepath):
    import pytesseract
    from PIL import Image

    tesseract_cmd = current_app.config.get("TESSERACT_CMD")
    if tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

    image = Image.open(filepath)
    text = pytesseract.image_to_string(image).strip()
    if not text:
        raise ExtractionError("No text could be detected in the image.")
    return text


# ---------------------------------------------------------------------------
# Page/section range helpers — power the "which part of the document" picker
# on the Simplify Notes screen. Real PDFs use actual page boundaries; DOCX,
# TXT, and OCR text (which have no real pages) fall back to even character
# chunks acting as "pseudo pages" so the same UI works for every file type.
# ---------------------------------------------------------------------------

def has_page_markers(raw_text):
    return bool(PAGE_MARKER_RE.search(raw_text or ""))


def estimate_page_count(raw_text):
    if not raw_text:
        return 0
    if has_page_markers(raw_text):
        return len(PAGE_MARKER_RE.findall(raw_text))
    return max(1, -(-len(raw_text) // CHARS_PER_PSEUDO_PAGE))  # ceil division


def get_text_for_page_range(raw_text, start_page, end_page):
    """Return text for pages [start_page, end_page], 1-indexed and inclusive.
    Uses real [[PAGE:n]] markers for PDFs; falls back to even character
    chunks for DOCX/TXT/OCR text that has no real page concept."""
    if not raw_text:
        return ""
    start_page = max(1, int(start_page))
    end_page = max(start_page, int(end_page))

    if has_page_markers(raw_text):
        matches = list(PAGE_MARKER_RE.finditer(raw_text))
        selected = []
        for idx, m in enumerate(matches):
            page_num = idx + 1
            if start_page <= page_num <= end_page:
                content_start = m.end()
                content_end = matches[idx + 1].start() if idx + 1 < len(matches) else len(raw_text)
                selected.append(raw_text[content_start:content_end].strip())
        return "\n\n".join(selected)

    start_idx = (start_page - 1) * CHARS_PER_PSEUDO_PAGE
    end_idx = end_page * CHARS_PER_PSEUDO_PAGE
    return raw_text[start_idx:end_idx]
